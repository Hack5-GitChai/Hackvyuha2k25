import os
import streamlit as st
from docx import Document

from utils.file_handler import read_file
from utils.classifier import classify_lines
from utils.template_manager import load_templates
from utils.formatter import apply_formatting, generate_docx_from_blocks
from utils.toc_generator import generate_toc
from utils.previewer import generate_pdf_preview, convert_pdf_to_images, img_to_base64
from utils.equation_renderer import clean_equation_input, render_equation_to_image, insert_equation_into_doc
from flask import Flask, render_template, redirect, url_for

# -------------------- Streamlit Setup --------------------
st.set_page_config(page_title="DocuMorph AI", layout="wide")
st.title("📄 DocuMorph AI – Intelligent Document Formatter")

# -------------------- Upload Section --------------------
st.subheader("📤 Upload a .txt or .docx file")
uploaded_file = st.file_uploader("Choose a file", type=["txt", "docx"])

if uploaded_file:
    raw_lines = read_file(uploaded_file)
    template_options = load_templates()
    selected_template = list(template_options.keys())[0]  # Default template selection
    style_json = template_options[selected_template]

    # Pre-classify lines immediately
    classified_lines = [(line, classify_lines(line)) for line in raw_lines]

    # -------------------- Apply Formatting Immediately --------------------
    with st.spinner("🛠️ Formatting document..."):
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)
        output_docx_path = os.path.join(output_dir, "formatted_output.docx")

        final_doc = Document()
        final_doc.add_section()

        # TOC
        toc_paras = generate_toc(classified_lines)
        for para in toc_paras.paragraphs:
            new_para = final_doc.add_paragraph(para.text)
            new_para.style = para.style

        # Content + equations
        equation_counter = 1
        eq_style = style_json.get("equation", {
            "width_inch": 3.5,
            "caption": {},
            "alignment": "center"
        })

        for text, label in classified_lines:
            if text.strip().lower().startswith("equation:"):
                math_expr = text.split(":", 1)[1].strip()
                latex = clean_equation_input(math_expr)
                img_path = render_equation_to_image(latex)
                insert_equation_into_doc(final_doc, img_path, eq_style, equation_number=equation_counter)
                equation_counter += 1
            else:
                para = final_doc.add_paragraph(text)
                style_map = {
                    "heading": "Heading 1",
                    "subheading": "Heading 2",
                    "body": "Normal"
                }
                para.style = style_map.get(label.lower(), "Normal")

        final_doc.save(output_docx_path)

    pdf_path = generate_pdf_preview(output_docx_path)

    # -------------------- Scrollable Preview Window --------------------
    st.subheader("🖨️ Document Preview (Like Google Docs Print Preview)")
    if pdf_path and os.path.exists(pdf_path):
        images = convert_pdf_to_images(pdf_path)
        if images:
            st.markdown("""
                <style>
                .pdf-viewer-box {
                    max-height: 700px;
                    overflow-y: scroll;
                    background-color: #f9f9f9;
                    padding: 1rem;
                    border: 1px solid #ccc;
                }
                .pdf-page {
                    display: block;
                    margin: 0 auto 20px auto;
                    box-shadow: 0 0 6px rgba(0,0,0,0.2);
                    border-radius: 8px;
                    max-width: 100%;
                }
                </style>
                <div class="pdf-viewer-box">
            """, unsafe_allow_html=True)

            for img in images:
                st.markdown(f'<img class="pdf-page" src="data:image/png;base64,{img_to_base64(img)}"/>', unsafe_allow_html=True)

            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.warning("⚠️ Could not render PDF preview.")
    else:
        st.error("❌ PDF generation failed.")

    # -------------------- Editable Classification --------------------
    st.subheader("✏️ Optional: Edit Detected Structure")
    updated_lines = []
    for i, (text, label) in enumerate(classified_lines):
        col1, col2 = st.columns([4, 2])
        with col1:
            st.text(text)
        with col2:
            new_label = st.selectbox(
                "Label", ["heading", "subheading", "body"],
                index=["heading", "subheading", "body"].index(label),
                key=f"label_{i}"
            )
        updated_lines.append((text, new_label))

    # -------------------- Final Export --------------------
    st.subheader("📥 Export Formatted Document")
    if st.button("Apply Edits & Export"):
        final_doc = Document()
        final_doc.add_section()

        toc_paras = generate_toc(updated_lines)
        for para in toc_paras.paragraphs:
            new_para = final_doc.add_paragraph(para.text)
            new_para.style = para.style

        equation_counter = 1
        for text, label in updated_lines:
            if text.strip().lower().startswith("equation:"):
                math_expr = text.split(":", 1)[1].strip()
                latex = clean_equation_input(math_expr)
                img_path = render_equation_to_image(latex)
                insert_equation_into_doc(final_doc, img_path, eq_style, equation_number=equation_counter)
                equation_counter += 1
            else:
                para = final_doc.add_paragraph(text)
                para.style = style_map.get(label.lower(), "Normal")

        final_doc.save(output_docx_path)
        pdf_path = generate_pdf_preview(output_docx_path)
        st.success("✅ Document updated and saved!")

    col1, col2 = st.columns(2)
    with col1:
        with open(output_docx_path, "rb") as f:
            st.download_button("⬇ Download DOCX", f, file_name="formatted_output.docx")
    with col2:
        if pdf_path and os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                st.download_button("⬇ Download PDF", f, file_name="formatted_output.pdf")
